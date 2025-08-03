"""Сервис парсинга файлов"""

import os
from abc import ABC, abstractmethod
from typing import List, Optional, Dict, Any, Tuple, Set
from pathlib import Path

import PyPDF2
from docx import Document as DocxDocument

from app.utils.logging_config import get_logger
from app.utils.helpers import get_file_hash, get_file_size_mb


logger = get_logger(__name__)


class FileParserError(Exception):
    """Исключение для ошибок парсинга файлов"""
    pass


class BaseFileParser(ABC):
    """Базовый класс для парсеров файлов"""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """
        Извлечь текст из файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Извлечённый текст
            
        Raises:
            FileParserError: При ошибке парсинга
        """
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """
        Получить список поддерживаемых расширений
        
        Returns:
            Список расширений файлов
        """
        pass


class TxtFileParser(BaseFileParser):
    """Парсер для TXT файлов"""
    
    def extract_text(self, file_path: str) -> str:
        """Извлечь текст из TXT файла"""
        try:
            # Попробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'windows-1251', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        logger.info(f"Файл {file_path} успешно прочитан с кодировкой {encoding}")
                        return content
                except UnicodeDecodeError:
                    continue
            
            # Если ни одна кодировка не подошла, используем режим с игнорированием ошибок
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                logger.warning(f"Файл {file_path} прочитан с игнорированием ошибок кодировки")
                return content
                
        except Exception as e:
            logger.error(f"Ошибка чтения TXT файла {file_path}: {e}")
            raise FileParserError(f"Не удалось прочитать TXT файл: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['txt']


class PdfFileParser(BaseFileParser):
    """Парсер для PDF файлов"""
    
    def extract_text(self, file_path: str) -> str:
        """Извлечь текст из PDF файла"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Проверить, что PDF не защищён паролем
                if pdf_reader.is_encrypted:
                    logger.error(f"PDF файл {file_path} защищён паролем")
                    raise FileParserError("PDF файл защищён паролем")
                
                # Извлечь текст из всех страниц
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():  # Проверить, что страница не пустая
                        text_content.append(page_text)
                        logger.debug(f"Извлечён текст со страницы {page_num + 1}")
            
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                logger.warning(f"PDF файл {file_path} не содержит извлекаемого текста")
                raise FileParserError("PDF файл не содержит текста или текст не может быть извлечён")
            
            logger.info(f"Успешно извлечён текст из PDF файла {file_path}")
            return full_text
            
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"Ошибка чтения PDF файла {file_path}: {e}")
            raise FileParserError(f"Не удалось прочитать PDF файл: {e}")
        except Exception as e:
            logger.error(f"Общая ошибка при парсинге PDF {file_path}: {e}")
            raise FileParserError(f"Ошибка при обработке PDF файла: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['pdf']


class DocxFileParser(BaseFileParser):
    """Парсер для DOCX файлов"""
    
    def extract_text(self, file_path: str) -> str:
        """Извлечь текст из DOCX файла"""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Извлечь текст из всех параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Извлечь текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                logger.warning(f"DOCX файл {file_path} не содержит текста")
                raise FileParserError("DOCX файл не содержит текста")
            
            logger.info(f"Успешно извлечён текст из DOCX файла {file_path}")
            return full_text
            
        except Exception as e:
            logger.error(f"Ошибка чтения DOCX файла {file_path}: {e}")
            raise FileParserError(f"Не удалось прочитать DOCX файл: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['docx']


class FileParserService:
    """Основной сервис парсинга файлов"""
    
    def __init__(self):
        # Регистрация парсеров
        self.parsers = {
            'txt': TxtFileParser(),
            'pdf': PdfFileParser(),
            'docx': DocxFileParser()
        }
        
        logger.info("FileParserService инициализирован")
    
    def get_supported_extensions(self) -> List[str]:
        """Получить список всех поддерживаемых расширений"""
        extensions = []
        for parser in self.parsers.values():
            extensions.extend(parser.get_supported_extensions())
        return extensions
    
    def is_supported_file(self, filename: str) -> bool:
        """
        Проверить, поддерживается ли файл
        
        Args:
            filename: Имя файла
            
        Returns:
            True если файл поддерживается
        """
        file_extension = filename.split('.')[-1].lower()
        return file_extension in self.parsers
    
    def parse_file(self, file_path: str) -> dict:
        """
        Парсить файл и извлечь текст
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Словарь с результатами парсинга
            
        Raises:
            FileParserError: При ошибке парсинга
        """
        if not os.path.exists(file_path):
            raise FileParserError(f"Файл не найден: {file_path}")
        
        # Определить тип файла
        filename = os.path.basename(file_path)
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension not in self.parsers:
            raise FileParserError(f"Неподдерживаемый тип файла: {file_extension}")
        
        # Получить парсер
        parser = self.parsers[file_extension]
        
        try:
            # Извлечь текст
            logger.info(f"Начинаю парсинг файла {file_path}")
            extracted_text = parser.extract_text(file_path)
            
            # Собрать метаданные
            file_stats = {
                'filename': filename,
                'file_path': file_path,
                'file_type': file_extension,
                'file_size_mb': get_file_size_mb(file_path),
                'file_hash': get_file_hash(file_path),
                'extracted_text': extracted_text,
                'text_length': len(extracted_text),
                'word_count': len(extracted_text.split()),
                'line_count': len(extracted_text.split('\n'))
            }
            
            logger.info(
                f"Файл {filename} успешно обработан. "
                f"Извлечено {file_stats['text_length']} символов, "
                f"{file_stats['word_count']} слов"
            )
            
            return file_stats
            
        except Exception as e:
            logger.error(f"Ошибка при парсинге файла {file_path}: {e}")
            raise FileParserError(f"Ошибка парсинга: {e}")
    
    def validate_file(self, file_path: str, max_size_mb: Optional[float] = None) -> bool:
        """
        Валидировать файл перед парсингом
        
        Args:
            file_path: Путь к файлу
            max_size_mb: Максимальный размер файла в MB
            
        Returns:
            True если файл валиден
            
        Raises:
            FileParserError: При ошибке валидации
        """
        if not os.path.exists(file_path):
            raise FileParserError(f"Файл не найден: {file_path}")
        
        # Проверить размер файла
        if max_size_mb:
            file_size = get_file_size_mb(file_path)
            if file_size > max_size_mb:
                raise FileParserError(
                    f"Файл слишком большой: {file_size:.2f}MB > {max_size_mb}MB"
                )
        
        # Проверить расширение
        if not self.is_supported_file(os.path.basename(file_path)):
            filename = os.path.basename(file_path)
            file_extension = filename.split('.')[-1].lower()
            raise FileParserError(f"Неподдерживаемый тип файла: {file_extension}")
        
        return True


# Глобальный экземпляр сервиса
file_parser_service = FileParserService() 